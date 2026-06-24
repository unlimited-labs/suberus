"""Generate the Suberus conference abstract DOCX template + a filled example.

A full professional extended-abstract layout: bold centred title,
superscript-marked authors, italic affiliations, Keywords line, numbered
sections, a figure (with caption), a table (with caption), and a numbered
References list with in-text citations.

The header order also matches Suberus's DOCX extractor, so an uploaded file
round-trips into title / authors / affiliations / keywords automatically:

    Title -> Authors (superscript) -> Affiliations (leading marker + institution)
    -> E-mails -> Keywords: -> "Abstract" heading + body/sections/figures/refs

Run:  py template/build_template.py
"""

import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from PIL import Image, ImageDraw, ImageFont

HINT = RGBColor(0x80, 0x80, 0x80)
TITLE_PT, AUTHOR_PT, AFF_PT, BODY_PT, CAP_PT = 14, 12, 9, 11, 9
C, J = WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY


# --- low-level helpers -----------------------------------------------------

def base(doc):
    s = doc.sections[0]
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(s, m, Cm(2.5))
    n = doc.styles["Normal"]
    n.font.name = "Times New Roman"
    n.font.size = Pt(BODY_PT)
    n.paragraph_format.space_after = Pt(0)
    n.paragraph_format.line_spacing = 1.0


def run(p, text, *, size=None, bold=False, italic=False, sup=False, color=None):
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    if size:
        r.font.size = Pt(size)
    r.bold, r.italic, r.font.superscript = bold, italic, sup
    if color:
        r.font.color.rgb = color
    return r


def para(doc, align=None, before=0, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def heading(doc, text):
    p = para(doc, before=8, after=3)
    run(p, text, bold=True)
    return p


# --- figure images (drawn with PIL, embedded via BytesIO) ------------------

def example_figure():
    """Recrystallized fraction vs inter-pass time: labelled axes + units."""
    import math
    W, H = 1000, 620
    T_MAX, X_MAX = 10, 100  # x: time (s) 0-10, y: fraction (%) 0-100
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    ox, oy, w, h = 150, 45, 780, 470  # plot box
    try:
        ax = ImageFont.truetype("times.ttf", 28)
        tk = ImageFont.truetype("times.ttf", 22)
    except OSError:
        ax = tk = ImageFont.load_default()

    def ctext(xy, text, font, fill, anchor):
        l, t, r, b = d.textbbox((0, 0), text, font=font)
        tw, th = r - l, b - t
        x, y = xy
        if "m" in anchor:    # h-centre
            x -= tw / 2 + l
        if "r" in anchor:    # right-align
            x -= tw + l
        if "M" in anchor:    # v-centre
            y -= th / 2 + t
        d.text((x, y), text, fill=fill, font=font)

    d.rectangle([ox, oy, ox + w, oy + h], outline="black", width=2)
    for i in range(0, 6):  # x ticks (0,2,4,6,8,10 s)
        x = ox + w * i / 5
        d.line([x, oy, x, oy + h], fill="#e6e6e6")
        d.line([x, oy + h, x, oy + h + 8], fill="black", width=2)
        ctext((x, oy + h + 12), str(T_MAX * i // 5), tk, "black", "m")
    for i in range(0, 6):  # y ticks (0,20,...,100 %)
        y = oy + h - h * i / 5
        d.line([ox, y, ox + w, y], fill="#e6e6e6")
        d.line([ox - 8, y, ox, y], fill="black", width=2)
        ctext((ox - 14, y), str(X_MAX * i // 5), tk, "black", "rM")

    def curve(k, color):
        pts = []
        for j in range(201):
            t = T_MAX * j / 200
            frac = 1 - math.exp(-k * ((t / T_MAX) ** 1.6))
            pts.append((ox + (t / T_MAX) * w, oy + h - frac * h))
        d.line(pts, fill=color, width=4)

    curve(3.2, "#1f4fd8")   # 1000 C
    curve(6.5, "#d81f1f")   # 1100 C

    # axis titles (with units); y title drawn rotated
    ctext((ox + w / 2, oy + h + 46), "Inter-pass time, t (s)", ax, "black", "m")
    yl = Image.new("RGBA", (520, 50), (0, 0, 0, 0))
    ImageDraw.Draw(yl).text((0, 0), "Recrystallized fraction, X (%)",
                            font=ax, fill="black")
    yl = yl.crop(yl.getbbox()).rotate(90, expand=True)
    img.paste(yl, (28, int(oy + h / 2 - yl.height / 2)), yl)

    # legend (top-left, where the curves are low)
    lx, ly = ox + 30, oy + 30
    for dy, color, lab in ((0, "#d81f1f", "1100 °C"),
                           (38, "#1f4fd8", "1000 °C")):
        d.line([lx, ly + dy + 12, lx + 46, ly + dy + 12], fill=color, width=4)
        ctext((lx + 58, ly + dy + 12), lab, tk, "black", "M")

    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def placeholder_figure():
    W, H = 900, 480
    img = Image.new("RGB", (W, H), "#f4f4f4")
    d = ImageDraw.Draw(img)
    d.rectangle([4, 4, W - 5, H - 5], outline="#9aa0a6", width=3)
    d.line([4, 4, W - 5, H - 5], fill="#cfcfcf", width=2)
    d.line([W - 5, 4, 4, H - 5], fill="#cfcfcf", width=2)
    try:
        fnt = ImageFont.truetype("arial.ttf", 34)
    except OSError:
        fnt = ImageFont.load_default()
    text = "Insert figure / chart here"
    l, t, r, b = d.textbbox((0, 0), text, font=fnt)
    d.text(((W - (r - l)) / 2 - l, (H - (b - t)) / 2 - t), text, fill="#666666", font=fnt)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


def add_figure(doc, buf, caption_runs):
    p = para(doc, C, before=6, after=2)
    p.add_run().add_picture(buf, width=Cm(11))
    cap = para(doc, C, after=8)
    for text, kw in caption_runs:
        run(cap, text, size=CAP_PT, **kw)


def add_table(doc, headers, rows, caption_runs):
    cap = para(doc, C, before=6, after=2)
    for text, kw in caption_runs:
        run(cap, text, size=CAP_PT, **kw)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, htext in zip(t.rows[0].cells, headers):
        c.paragraphs[0].alignment = C
        run(c.paragraphs[0], htext, size=AFF_PT, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.paragraphs[0].alignment = C
            run(c.paragraphs[0], v, size=AFF_PT)
    para(doc, after=4)  # spacer after table


# --- document body ---------------------------------------------------------

def header(doc, example):
    # 1. Title
    t = para(doc, C, after=10)
    if example:
        run(t, "Microstructure Evolution in Hot-Rolled Steel During "
                "Multi-Pass Deformation", size=TITLE_PT, bold=True)
    else:
        run(t, "Title of the Abstract", size=TITLE_PT, bold=True)
        run(t, "  (Times New Roman 14 pt, bold, centred; max 2 lines)",
            size=TITLE_PT, bold=True, color=HINT)

    # 2. Authors (superscript markers; * = presenting)
    a = para(doc, C, after=2)
    if example:
        for i, (name, mark) in enumerate(
            [("Anna Kowalska", "1,*"), ("Jan Nowak", "2"),
             ("Maria Lewandowska", "1")]):
            if i:
                run(a, ", ", size=AUTHOR_PT)
            run(a, name, size=AUTHOR_PT)
            run(a, mark, size=AUTHOR_PT, sup=True)
    else:
        run(a, "Given Name Surname", size=AUTHOR_PT)
        run(a, "1,*", size=AUTHOR_PT, sup=True)
        run(a, ", Given Name Surname", size=AUTHOR_PT)
        run(a, "2", size=AUTHOR_PT, sup=True)
        run(a, "   (presenting author marked with *)", size=AUTHOR_PT, color=HINT)

    # 3. Affiliations (each line: marker + institution)
    affs = ([("1", "AGH University of Krakow, Faculty of Metals Engineering, "
                   "al. Mickiewicza 30, 30-059 Krakow, Poland"),
             ("2", "Lukasiewicz Research Network, Institute for Ferrous "
                   "Metallurgy, K. Miarki 12, 44-100 Gliwice, Poland")]
            if example else
            [("1", "Department, Institution, City, Country"),
             ("2", "Department, Institution, City, Country")])
    for i, (mark, text) in enumerate(affs):
        p = para(doc, C, after=(6 if i == len(affs) - 1 else 1))
        run(p, mark, size=AFF_PT, sup=True, italic=True)
        run(p, " " + text, size=AFF_PT, italic=True)

    # 4. E-mails (author order)
    e = para(doc, C, after=10)
    if example:
        run(e, "E-mails: a.kowalska@agh.edu.pl, j.nowak@imz.pl, "
               "m.lewandowska@agh.edu.pl", size=AFF_PT, italic=True)
    else:
        run(e, "E-mails: ", size=AFF_PT, italic=True)
        run(e, "presenting@author.org, co.author@institution.org "
               "(in author order)", size=AFF_PT, italic=True, color=HINT)

    # 5. Keywords
    k = para(doc, J, after=10)
    run(k, "Keywords: ", bold=True)
    if example:
        run(k, "hot rolling, microstructure, recrystallization, "
               "finite element modelling, steel")
    else:
        run(k, "keyword one, keyword two, keyword three", color=HINT)
        run(k, "  (3-6, comma-separated)", color=HINT)


def build(doc, example):
    base(doc)
    header(doc, example)

    # The whole body IS the abstract: continuous prose (no paper sections),
    # with one table, one figure and a references list. The paragraph after the
    # Keywords line is already the BODY zone for the extractor (no heading
    # needed). Citations are bracketed [1], [2] in order of appearance.
    p = para(doc, J, after=4)
    if example:
        run(p, "Grain refinement during thermomechanical processing controls "
               "the strength and toughness of structural steels [1], yet "
               "industrial multi-pass schedules retain strain between passes in "
               "a way that single-pass Johnson-Mehl-Avrami-Kolmogorov (JMAK) "
               "calibrations do not capture [2]. This work quantifies how "
               "inter-pass time and temperature govern recrystallization in "
               "low-carbon steel. Plane-strain compression specimens were "
               "deformed between 900 and 1100 °C following the four-pass "
               "schedule of Table 1, then interrupted-quenched for "
               "metallography.")
    else:
        run(p, "Write your abstract as continuous text (Times New Roman 11 pt, "
               "justified, single spacing). State the problem, the method, the "
               "key results and the conclusion in one flowing narrative - no "
               "section headings. You may include one table and one figure, and "
               "cite references in order as [1], [2]. Keep the whole document "
               "to two pages.", color=HINT)

    if example:
        add_table(doc,
                  ["Pass", "Temperature (°C)", "Strain, ε (–)",
                   "Inter-pass time (s)"],
                  [["1", "1100", "0.30", "10"], ["2", "1050", "0.30", "5"],
                   ["3", "1000", "0.25", "3"], ["4", "950", "0.25", "–"]],
                  [("Table 1. ", dict(bold=True)),
                   ("Four-pass plane-strain compression schedule.", {})])
    else:
        add_table(doc, ["Column", "Column", "Column"],
                  [["", "", ""], ["", "", ""]],
                  [("Table 1. ", dict(bold=True)),
                   ("Caption above the table (Times New Roman 9 pt, centred).",
                    dict(color=HINT))])

    p = para(doc, J, after=4)
    if example:
        run(p, "Figure 1 shows the recrystallized fraction as a function of "
               "inter-pass time at two temperatures. Above 1000 °C the kinetics "
               "are dominated by inter-pass time rather than peak strain; a "
               "JMAK law fitted to single-pass data over-predicts grain "
               "refinement in the multi-pass case by up to 18% [2].")
    else:
        run(p, "Continue the narrative: present and interpret the results and "
               "refer to the figure as Figure 1.", color=HINT)

    if example:
        add_figure(doc, example_figure(),
                   [("Fig. 1. ", dict(bold=True)),
                    ("Recrystallized fraction X(t) versus inter-pass time at "
                     "1000 °C and 1100 °C.", {})])
    else:
        add_figure(doc, placeholder_figure(),
                   [("Fig. 1. ", dict(bold=True)),
                    ("Caption below the figure (Times New Roman 9 pt, centred).",
                     dict(color=HINT))])

    p = para(doc, J)
    if example:
        run(p, "A retained-strain correction to the JMAK law reconciles "
               "laboratory and industrial bar-mill data, improving grain-size "
               "prediction across the schedule [3].")
    else:
        run(p, "Close with the conclusion in the same paragraph flow.",
            color=HINT)

    # References (the only heading; not a paper section)
    heading(doc, "References")
    refs = ([
        "F. J. Humphreys, M. Hatherly, Recrystallization and Related "
        "Annealing Phenomena, 2nd ed., Elsevier, Oxford, 2004.",
        "C. M. Sellars, J. A. Whiteman, Recrystallization and grain growth in "
        "hot rolling, Met. Sci. 13 (1979) 187-194.",
        "A. Kowalska, J. Nowak, Retained strain in multi-pass hot rolling, "
        "Mater. Sci. Eng. A 812 (2021) 141-152.",
    ] if example else [
        "Author A. A., Author B. B., Title of the work, Journal Vol. (Year) "
        "pages.",
        "Author C. C., Title of the book, Publisher, City, Year.",
    ])
    for i, r in enumerate(refs, 1):
        rp = para(doc, J, after=2)
        rp.paragraph_format.left_indent = Cm(0.8)
        rp.paragraph_format.first_line_indent = Cm(-0.8)
        run(rp, f"[{i}] ", size=AFF_PT)
        run(rp, r, size=AFF_PT, color=(None if example else HINT))

    # Footer note
    f = para(doc, before=10, after=0)
    run(f, "Format: A4, 2.5 cm margins, Times New Roman. Title 14 pt bold; "
           "authors 12 pt; affiliations/e-mails/captions/references 9 pt; body "
           "11 pt justified. Table caption above, figure caption below. "
           "Two pages maximum.", size=8, italic=True, color=HINT)


for fname, ex in (("abstract-template.docx", False), ("abstract-example.docx", True)):
    doc = Document()
    build(doc, ex)
    doc.save("template/docx/" + fname)
    print("wrote template/docx/" + fname)
