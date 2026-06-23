"""Generate the representative .docx baseline fixtures (run inside the docx-api
test image: `python test_fixtures/generate_fixtures.py`).

Each document is deliberately small, uses generic / lorem-ipsum text only, and
focuses on a cluster of Word features so the normalize pipeline (Pandoc +
LibreOffice + stylecss + math) has a broad, reproducible surface to pin.

The generated .docx are committed; this script lets us regenerate/extend them.
Goldens are keyed off the *normalized HTML*, not the .docx bytes, so the only
reason to keep generation deterministic is reviewable git diffs — we therefore
pin core-property timestamps.

Features that cannot be authored programmatically (autoshapes, text boxes,
EMF/WMF vector images, real MathType/Word-equation authoring) are supplied by a
human under test_fixtures/real/ — see README.
"""

import io
import zipfile
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

HERE = Path(__file__).parent
DOCS = HERE / "docs"
MATHTYPE_BIN = HERE / "mathtype_equation_dsmt4.bin"

# Fixed timestamp so regeneration yields reviewable diffs, not timestamp churn.
FIXED_TS = datetime(2024, 1, 1, 0, 0, 0, tzinfo=timezone.utc)

LOREM = (
    "Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod "
    "tempor incididunt ut labore et dolore magna aliqua."
)

M = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _stamp(doc: Document) -> None:
    doc.core_properties.created = FIXED_TS
    doc.core_properties.modified = FIXED_TS


def _save(doc: Document, name: str) -> None:
    _stamp(doc)
    DOCS.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOCS / name))
    print(f"  wrote docs/{name}")


def gen_text_formatting() -> None:
    doc = Document()
    doc.add_heading("Text formatting", level=1)

    for label, align in [
        ("Left aligned. " + LOREM, WD_ALIGN_PARAGRAPH.LEFT),
        ("Right aligned. " + LOREM, WD_ALIGN_PARAGRAPH.RIGHT),
        ("Centered. " + LOREM, WD_ALIGN_PARAGRAPH.CENTER),
        ("Justified. " + LOREM + " " + LOREM, WD_ALIGN_PARAGRAPH.JUSTIFY),
    ]:
        p = doc.add_paragraph(label)
        p.alignment = align

    sizes = doc.add_paragraph()
    for pt in (8, 11, 14, 24):
        r = sizes.add_run(f"{pt}pt ")
        r.font.size = Pt(pt)

    decorated = doc.add_paragraph()
    decorated.add_run("bold ").bold = True
    decorated.add_run("italic ").italic = True
    decorated.add_run("underline ").underline = True
    caps = decorated.add_run("smallcaps ")
    caps.font.all_caps = True
    colored = decorated.add_run("red")
    colored.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    spaced = doc.add_paragraph("Paragraph with extra spacing before and after.")
    spaced.paragraph_format.space_before = Pt(18)
    spaced.paragraph_format.space_after = Pt(18)

    _save(doc, "text-formatting.docx")


def _add_paren_numbering(doc: Document) -> tuple[int, int]:
    """Append two paren-delimited ("a)" / "i)") abstract numberings to the
    document's numbering part. Returns the (lowerLetter, lowerRoman) numIds."""
    numbering = doc.part.numbering_part.numbering_definitions._numbering

    def abstract(abs_id: int, fmt: str) -> str:
        return (
            f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abs_id}">'
            f'<w:lvl w:ilvl="0"><w:start w:val="1"/>'
            f'<w:numFmt w:val="{fmt}"/><w:lvlText w:val="%1)"/>'
            f'<w:lvlJc w:val="left"/>'
            f'<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr>'
            f"</w:lvl></w:abstractNum>"
        )

    def num(num_id: int, abs_id: int) -> str:
        return (
            f'<w:num {nsdecls("w")} w:numId="{num_id}">'
            f'<w:abstractNumId w:val="{abs_id}"/></w:num>'
        )

    # abstractNum elements must precede num elements in the schema.
    first_num = numbering.find(qn("w:num"))
    for abs_id, fmt in ((100, "lowerLetter"), (101, "lowerRoman")):
        el = parse_xml(abstract(abs_id, fmt))
        if first_num is not None:
            first_num.addprevious(el)
        else:
            numbering.append(el)
    for num_id, abs_id in ((100, 100), (101, 101)):
        numbering.append(parse_xml(num(num_id, abs_id)))
    return 100, 101


def _numbered(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph(text, style="List Paragraph")
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numPr.append(ilvl)
    numPr.append(nid)
    pPr.append(numPr)


def gen_lists() -> None:
    doc = Document()
    doc.add_heading("Lists", level=1)

    doc.add_paragraph("Bulleted:")
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Nested bullet", style="List Bullet 2")
    doc.add_paragraph("Second bullet", style="List Bullet")

    doc.add_paragraph("Decimal numbered:")
    doc.add_paragraph("First numbered", style="List Number")
    doc.add_paragraph("Second numbered", style="List Number")

    letter_id, roman_id = _add_paren_numbering(doc)
    doc.add_paragraph("Lettered with paren:")
    _numbered(doc, "Alpha item", letter_id)
    _numbered(doc, "Beta item", letter_id)
    doc.add_paragraph("Roman with paren:")
    _numbered(doc, "First roman", roman_id)
    _numbered(doc, "Second roman", roman_id)

    _save(doc, "lists.docx")


def _mark_header_row(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")} w:val="true"/>'))


def gen_tables() -> None:
    doc = Document()
    doc.add_heading("Tables", level=1)

    t = doc.add_table(rows=3, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["Header A", "Header B", "Header C"]
    for c, text in enumerate(headers):
        t.cell(0, c).text = text
    _mark_header_row(t.rows[0])
    for r in range(1, 3):
        for c in range(3):
            t.cell(r, c).text = f"r{r}c{c}"
    centered = t.cell(1, 1)
    centered.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph("Merged cells:")
    m = doc.add_table(rows=2, cols=3)
    m.style = "Table Grid"
    m.cell(0, 0).text = "spans two columns"
    m.cell(0, 0).merge(m.cell(0, 1))  # colspan
    m.cell(0, 2).text = "tall"
    m.cell(0, 2).merge(m.cell(1, 2))  # rowspan
    m.cell(1, 0).text = "a"
    m.cell(1, 1).text = "b"

    _save(doc, "tables.docx")


def _png_bytes(color: tuple[int, int, int], size=(120, 80)) -> io.BytesIO:
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="PNG")
    buf.seek(0)
    return buf


def _jpg_bytes(color: tuple[int, int, int], size=(100, 100)) -> io.BytesIO:
    buf = io.BytesIO()
    Image.new("RGB", size, color).save(buf, format="JPEG", quality=90)
    buf.seek(0)
    return buf


def gen_images() -> None:
    doc = Document()
    doc.add_heading("Images", level=1)

    doc.add_paragraph("A PNG (passthrough), sized to 1.5in:")
    doc.add_picture(_png_bytes((40, 110, 200)), width=Inches(1.5))

    doc.add_paragraph("A JPEG (rasterized to PNG), centered:")
    doc.add_picture(_jpg_bytes((200, 90, 40)), width=Inches(1.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    _save(doc, "images.docx")


def _omath_run(text: str) -> str:
    return f"<m:r><m:t>{text}</m:t></m:r>"


def _ssup(base: str, sup: str) -> str:
    return f"<m:sSup><m:e>{base}</m:e><m:sup>{sup}</m:sup></m:sSup>"


def _frac(num: str, den: str) -> str:
    return f"<m:f><m:num>{num}</m:num><m:den>{den}</m:den></m:f>"


def _add_equation(doc: Document, omath_body: str, display: bool) -> None:
    p = doc.add_paragraph()
    omath = f'<m:oMath xmlns:m="{M}">{omath_body}</m:oMath>'
    if display:
        xml = f'<m:oMathPara xmlns:m="{M}">{omath}</m:oMathPara>'
    else:
        xml = omath
    p._p.append(parse_xml(xml))


def gen_equations_omml() -> None:
    doc = Document()
    doc.add_heading("Native Word equations (OMML)", level=1)

    sqrt = (
        "<m:rad><m:deg/><m:e>"
        + _ssup(_omath_run("b"), _omath_run("2"))
        + _omath_run("-4ac")
        + "</m:e></m:rad>"
    )
    quad = _frac(_omath_run("-b±") + sqrt, _omath_run("2a"))
    doc.add_paragraph("Quadratic formula (display):")
    _add_equation(doc, _omath_run("x=") + quad, display=True)

    doc.add_paragraph("Fraction (inline):")
    _add_equation(doc, _frac(_omath_run("g"), _omath_run("4")), display=False)

    nary = (
        '<m:nary><m:naryPr><m:chr m:val="∫"/></m:naryPr>'
        f"<m:sub>{_omath_run('a')}</m:sub><m:sup>{_omath_run('b')}</m:sup>"
        f"<m:e>{_omath_run('f(x)dx')}</m:e></m:nary>"
    )
    doc.add_paragraph("Definite integral (display):")
    _add_equation(doc, nary, display=True)

    greek = _ssup(_omath_run("α"), _omath_run("2")) + _omath_run("+β")
    doc.add_paragraph("Greek with superscript (inline):")
    _add_equation(doc, greek, display=False)

    _save(doc, "equations-omml.docx")


def gen_mathtype() -> None:
    """Build a full, valid DOCX (so pandoc accepts it) embedding the real
    Equation.DSMT4 OLE object on its own paragraph."""
    ole = MATHTYPE_BIN.read_bytes()
    w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    o = "urn:schemas-microsoft-com:office:office"
    r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    content_types = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Default Extension="bin" ContentType="application/vnd.openxmlformats-officedocument.oleObject"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        "</Types>"
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rIdDoc" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/></Relationships>'
    )
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:document xmlns:w="{w}" xmlns:o="{o}" xmlns:r="{r}"><w:body>'
        '<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>'
        "<w:r><w:t>MathType equation</w:t></w:r></w:p>"
        "<w:p><w:r><w:t>An embedded MathType equation on its own line:</w:t></w:r></w:p>"
        "<w:p><w:r><w:object>"
        '<o:OLEObject Type="Embed" ProgID="Equation.DSMT4" r:id="rId1"/>'
        "</w:object></w:r></w:p>"
        "</w:body></w:document>"
    )
    doc_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject" '
        'Target="embeddings/oleObject1.bin"/></Relationships>'
    )

    DOCS.mkdir(parents=True, exist_ok=True)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("[Content_Types].xml", content_types)
        z.writestr("_rels/.rels", root_rels)
        z.writestr("word/document.xml", document)
        z.writestr("word/_rels/document.xml.rels", doc_rels)
        z.writestr("word/embeddings/oleObject1.bin", ole)
    (DOCS / "mathtype.docx").write_bytes(buf.getvalue())
    print("  wrote docs/mathtype.docx")


def gen_headings() -> None:
    doc = Document()
    doc.add_paragraph("Generic Paper Title", style="Title")
    doc.add_paragraph("A representative subtitle", style="Subtitle")
    doc.add_heading("Section one", level=1)
    doc.add_paragraph(LOREM)
    doc.add_heading("Subsection", level=2)
    doc.add_paragraph(LOREM)
    doc.add_heading("Sub-subsection", level=3)
    doc.add_paragraph(LOREM)
    _save(doc, "headings.docx")


def _kitchen(doc: Document, *, word: str, gamma_first: bool, exp: str, cell: str) -> None:
    doc.add_heading("Kitchen sink", level=1)
    doc.add_paragraph(f"The quick {word} brown fox jumps over the lazy dog.")

    blocks = ["Alpha block paragraph.", "Beta block paragraph.", "Gamma block paragraph."]
    if gamma_first:
        blocks = [blocks[2], blocks[0], blocks[1]]
    for b in blocks:
        doc.add_paragraph(b)

    doc.add_paragraph("An inline equation:")
    _add_equation(doc, _ssup(_omath_run("x"), _omath_run(exp)), display=False)

    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = "key"
    t.cell(0, 1).text = "value"
    t.cell(1, 0).text = "status"
    t.cell(1, 1).text = cell


def gen_kitchen_sink() -> None:
    v1 = Document()
    _kitchen(v1, word="quick", gamma_first=False, exp="2", cell="draft")
    _save(v1, "kitchen-sink-v1.docx")

    v2 = Document()
    _kitchen(v2, word="slow", gamma_first=True, exp="3", cell="final")
    _save(v2, "kitchen-sink-v2.docx")


def main() -> None:
    print("Generating baseline .docx fixtures...")
    gen_text_formatting()
    gen_lists()
    gen_tables()
    gen_images()
    gen_equations_omml()
    gen_mathtype()
    gen_headings()
    gen_kitchen_sink()
    print("Done.")


if __name__ == "__main__":
    main()
